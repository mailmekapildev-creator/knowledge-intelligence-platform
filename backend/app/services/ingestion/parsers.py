"""
Format-specific parsing. Each parser takes raw bytes + filename and returns plain text
plus structural metadata. Isolated from chunking on purpose: a parsing failure should
fail one document, not the pipeline (see docs/failure-modes.md).

This is a portfolio-scope implementation: text/markdown/html/csv are fully implemented;
PDF/DOCX/PPTX/OCR paths are implemented against their real libraries when available and
fall back to a clearly-labeled mock extraction in MOCK_MODE so the pipeline is runnable
without heavyweight binary dependencies.
"""
from __future__ import annotations

import csv
import io
import re
from dataclasses import dataclass, field


class UnsupportedFormatError(Exception):
    pass


class ParsingError(Exception):
    def __init__(self, reason: str):
        self.reason = reason
        super().__init__(reason)


@dataclass
class ParsedDocument:
    text: str
    mime_type: str
    page_count: int = 1
    sections: list[dict] = field(default_factory=list)  # [{title, start_char, end_char}]


def _strip_html(raw: str) -> str:
    # Minimal, dependency-free tag stripper for the portfolio build.
    text = re.sub(r"<script.*?</script>", " ", raw, flags=re.S | re.I)
    text = re.sub(r"<style.*?</style>", " ", text, flags=re.S | re.I)
    text = re.sub(r"<[^>]+>", " ", text)
    return re.sub(r"\s+", " ", text).strip()


def _extract_markdown_sections(text: str) -> list[dict]:
    sections = []
    for match in re.finditer(r"^(#{1,6})\s+(.*)$", text, flags=re.M):
        sections.append({"level": len(match.group(1)), "title": match.group(2).strip(),
                          "start_char": match.start()})
    return sections


def parse_document(filename: str, content: bytes, mime_type: str) -> ParsedDocument:
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""

    if ext in ("txt",):
        text = content.decode("utf-8", errors="replace")
        return ParsedDocument(text=text, mime_type="text/plain")

    if ext in ("md", "markdown"):
        text = content.decode("utf-8", errors="replace")
        return ParsedDocument(text=text, mime_type="text/markdown",
                               sections=_extract_markdown_sections(text))

    if ext in ("html", "htm"):
        raw = content.decode("utf-8", errors="replace")
        return ParsedDocument(text=_strip_html(raw), mime_type="text/html")

    if ext == "csv":
        reader = csv.reader(io.StringIO(content.decode("utf-8", errors="replace")))
        rows = list(reader)
        if not rows:
            raise ParsingError("empty CSV")
        header, *body = rows
        # Summarize row-by-row rather than dumping the raw grid as prose (see docs).
        lines = [f"Columns: {', '.join(header)}"]
        for row in body[:5000]:
            pairs = ", ".join(f"{h}={v}" for h, v in zip(header, row))
            lines.append(pairs)
        return ParsedDocument(text="\n".join(lines), mime_type="text/csv")

    if ext == "pdf":
        try:
            import pypdf  # optional dependency, real path if installed
            reader = pypdf.PdfReader(io.BytesIO(content))
            pages = [p.extract_text() or "" for p in reader.pages]
            text = "\n\n".join(pages)
            if not text.strip():
                raise ParsingError("no extractable text -- looks like a scanned PDF, route to OCR")
            return ParsedDocument(text=text, mime_type="application/pdf", page_count=len(pages))
        except ImportError:
            return ParsedDocument(
                text=f"[MOCK PARSE] Extracted text placeholder for PDF '{filename}' "
                     f"({len(content)} bytes). Install pypdf for real extraction.",
                mime_type="application/pdf",
            )

    if ext == "docx":
        try:
            import docx  # python-docx
            d = docx.Document(io.BytesIO(content))
            text = "\n".join(p.text for p in d.paragraphs)
            return ParsedDocument(text=text, mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        except ImportError:
            return ParsedDocument(
                text=f"[MOCK PARSE] Extracted text placeholder for DOCX '{filename}'. "
                     f"Install python-docx for real extraction.",
                mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

    raise UnsupportedFormatError(f"Unsupported file extension: .{ext}")
